import json
import openai
import boto3
import datetime
import re
from io import BytesIO,StringIO
import jiter
from datetime import datetime, timedelta
import nbformat
from docx import Document
from nbformat.v4 import new_notebook, new_code_cell, new_markdown_cell
from docx.shared import Mm
import base64
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import PageTemplate, Frame
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from openai import OpenAI
import os
from dotenv import load_dotenv
load_dotenv()

S3_CLIENT = boto3.client(
    "s3",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name="us-east-1"
)
S3_BUCKET_NAME = "generated-assignments"

# OpenAI API Key
 # Replace with a valid key

client = OpenAI(api_key=os.getenv("OPEN_AI_API_KEY"))

class Processor:
    def cleaner(self,output={}):
        if '```json' in output['body']:
            new_data = output['body'][output['body'].find('```json')+9:]
        elif 'error' in output['body']:
            print("ERROR")
            new_data = None
        else:
            new_data = output['body']

        if not isinstance(new_data, str):
            new_data = str(new_data)

        cleaned_text = new_data.strip("'\"")
        cleaned_text = cleaned_text.replace("\\\\\\", "")
        cleaned_text = cleaned_text.replace("\n", "")
        cleaned_text = cleaned_text.replace("\\n", "")
        cleaned_text = cleaned_text.replace("\\\\n ", "")
        cleaned_text = cleaned_text.replace("```", "")
        cleaned_text = cleaned_text.replace("'", "")
        cleaned_text = cleaned_text.replace("\\", "")

        print(cleaned_text)

        try:
            byte_data = cleaned_text.encode('utf-8')
            final_output = jiter.from_json(byte_data, partial_mode=True)
            return final_output

        except Exception as e:
            print(f"JSON Parsing Failed: {e}")
            return {"error": "Failed to parse GPT output", "raw": cleaned_text}


    def generate_assignment(self,topic, beginner=0, intermediate=0, advanced=0):
        """Generate coding interview questions using OpenAI GPT-4o."""
        prompt = f"""
            You are a helpful assistant generating coding interview questions for the topic: **{topic}**.

            Please generate the following:
            - {beginner} **beginner** questions (easy, concept-based)
            - {intermediate} **intermediate** questions (problem-solving based)
            - {advanced} **advanced** questions (tricky, involving complex data structures and algorithms, include a HINT in each)

            ### Output Format:
            Return a JSON object with keys `beginner`, `intermediate`, and `advanced`.  
            Each key should map to a list of items in this format:

            ```json
            {{
            "beginner": [
                [1, "Beginner question 1"],
                [2, "Beginner question 2"]
            ],
            "intermediate": [
                [1, "Intermediate question 1"],
                [2, "Intermediate question 2"]
            ],
            "advanced": [
                [1, "Advanced question 1 (HINT: Add a helpful hint here)"],
                [2, "Advanced question 2 (HINT: Add another hint here)"]
            ]
            }}
            """

        try:

            response = client.chat.completions.create(
                model="gpt-4o-mini", 
                messages=[
                    {"role": "system", "content": "You are an AI that generates structured JSON responses."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.8,
                max_tokens=1024
            )
            result = response.model_dump()
            return result["choices"][0]["message"]["content"]
        
        except openai.APIError as e:
            return {"error": "API request failed", "details": str(e)}


    def upload_pdf_to_s3(self,file_buffer, filename, metadata=None):
        content_type_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".html": "text/html",
            ".ipynb": "application/x-ipynb+json"
        }
        
        ext = '.' + filename.split('.')[-1]
        content_type = content_type_map.get(ext, 'application/octet-stream')

        extra_args = {'ContentType': content_type}
        
        if metadata:
            # S3 metadata keys and values must be strings
            extra_args['Metadata'] = {str(k): str(v) for k, v in metadata.items()}
        
        try:
            file_buffer.seek(0)
            S3_CLIENT.upload_fileobj(
                file_buffer,
                S3_BUCKET_NAME,
                filename,
                ExtraArgs=extra_args
            )
            return f"https://{S3_BUCKET_NAME}.s3.amazonaws.com/{filename}"
        except Exception as e:
            return {"error": "S3 upload failed", "details": str(e)}


    def sanitize_filename(self,topic):
        """Sanitize topic name for a safe filename."""
        return re.sub(r"[^a-zA-Z0-9_]+", "_", topic.lower().strip())


    def add_watermark_and_border(self,canvas_obj, doc):
        width, height = A4

        # Watermark
        logo_path = "image.png"
        logo = ImageReader(logo_path)
        page_width, page_height = A4
        logo_width, logo_height = 200, 200
        x = (page_width - logo_width) / 2
        y = (page_height - logo_height) / 2

        canvas_obj.saveState()
        try:
            canvas_obj.setFillAlpha(0.1)  # transparency for newer ReportLab versions
        except AttributeError:
            pass  # fallback for older versions, no transparency
        
        canvas_obj.drawImage(logo, x, y, width=logo_width, height=logo_height, mask='auto')
        canvas_obj.restoreState()

        # Border
        canvas_obj.setStrokeColor(colors.grey)
        canvas_obj.setLineWidth(2)
        canvas_obj.rect(10, 10, width - 20, height - 20)


    def generate_pdf_from_json(self,data,topic):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = []

        styles.add(ParagraphStyle(name='TitleCenter', fontSize=20, alignment=TA_CENTER, textColor=colors.darkblue, spaceAfter=20))
        styles.add(ParagraphStyle(name='LevelHeading', fontSize=14, textColor=colors.HexColor("#1F618D"), spaceBefore=12, spaceAfter=8))
        styles.add(ParagraphStyle(name='Question', fontSize=11, textColor=colors.black, leftIndent=10, spaceAfter=6))

        elements.append(Paragraph(f"<u>{topic}</u>", styles['TitleCenter']))
        elements.append(Spacer(1, 12))

        for level in ['beginner', 'intermediate', 'advanced']:
            elements.append(Paragraph(level.capitalize(), styles['LevelHeading']))
            for qid, question in data[level]:
                elements.append(Paragraph(f"<b>Q{qid}:</b> {question}", styles['Question']))
            elements.append(Spacer(1, 12))

        frame = Frame(20, 20, A4[0] - 40, A4[1] - 40, id='normal')
        template = PageTemplate(id='watermarked', frames=frame, onPage=self.add_watermark_and_border)
        doc.addPageTemplates([template])

        doc.build(elements)
        buffer.seek(0)
        return buffer


    def create_file(self,clean_data, file_type,topic):
        file_type = file_type.lower()
        with open("image.png", "rb") as img_file:
            image_bytes = img_file.read()
        
        if file_type == "pdf":
            buffer = self.generate_pdf_from_json(clean_data,topic)
            return buffer, "pdf"
        
        elif file_type == "docx":
            doc = Document()
            doc.add_picture(BytesIO(image_bytes), width=Mm(40))
            for level, questions in clean_data.items():
                doc.add_heading(level.title(), level=1)
                for q in questions:
                    doc.add_paragraph(f"{q[0]}. {q[1]}")
            buffer = BytesIO()
            doc.save(buffer)
            # save_locally(buffer, "debug_assignment.docx")

            return buffer, "docx"
        
        elif file_type == "html":
            encoded = base64.b64encode(image_bytes).decode("utf-8")
            img_tag = f"<img src='data:image/png;base64,{encoded}' style='display: block; width: 151px;' />"
            html_content = f"<html><body>{img_tag}"
            for level, questions in clean_data.items():
                html_content += f"<h2>{level.title()}</h2><ul>"
                for q in questions:
                    html_content += f"<li>{q[0]}. {q[1]}</li>"
                html_content += "</ul>"
            buffer = BytesIO(html_content.encode("utf-8"))
            # save_locally(buffer, "debug_assignment.html")
            return buffer,'html'
        
        elif file_type == "ipynb":
            

            nb = new_notebook()
            cells = []
            encoded = base64.b64encode(image_bytes).decode("utf-8")
            image_markdown = f"![image](data:image/png;base64,{encoded})"
            cells.append(new_markdown_cell(image_markdown))
            for level, questions in clean_data.items():
                cells.append(new_markdown_cell(f"## {level.capitalize()} Questions"))
                for q in questions:
                    question_num, question_text = q
                    cells.append(new_markdown_cell(f"**Q{question_num}.** {question_text}"))
                    cells.append(new_code_cell("# Your code here\n"))

            nb['cells'] = cells

            # Write to string buffer first
            string_buffer = StringIO()
            nbformat.write(nb, string_buffer)
            notebook_str = string_buffer.getvalue().encode("utf-8")  # Convert string to bytes

            buffer = BytesIO(notebook_str)  # Now you can use this for S3 upload or local save
            # save_locally(buffer, "debug_assignment.ipynb")
            return buffer,'ipynb'
            
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
    #For Debugging:
    # def save_locally(buffer, filename):
    #     debug_dir = "debug_outputs"
    #     os.makedirs(debug_dir, exist_ok=True)
    #     filepath = os.path.join(debug_dir, filename)
    #     with open(filepath, "wb") as f:
    #         buffer.seek(0)
    #         f.write(buffer.read())
    #     print(f"Saved file locally at: {filepath}")

class Gen:
    def generator(self,event, context=None):
        obj = Processor()
        try:
            tomorrow = datetime.now().date() + timedelta(days=1)
            tomorrow = tomorrow.strftime("%d%m%y")
            body = json.loads(event['request']["body"]) 
            topic = body.get("topic", "Python")
            beginner = int(body.get("beginner", 0))
            intermediate = int(body.get("intermediate", 0))
            advanced = int(body.get("advanced", 0))
            deadline = (body.get("deadline",tomorrow))
            file_type = body.get('file_type','pdf')
            batch = body.get('batch')

            # Generate assignment questions
            assignment_data = obj.generate_assignment(topic, beginner, intermediate, advanced)
            if isinstance(assignment_data, dict) and "error" in assignment_data:
                return {"statusCode": 500, "body": json.dumps(assignment_data)}

            clean_data = obj.cleaner({"body":json.dumps(assignment_data)})
            filename = f"assignment_{obj.sanitize_filename(topic)}_{deadline}.pdf"
            file_buffer, extension = obj.create_file(clean_data, file_type,topic)
            filename = f"assignment_{obj.sanitize_filename(topic)}_{deadline}.{extension}"
            metadata = {
                "batch_name": batch,
                "deadline": deadline
            }
            file_url = obj.upload_pdf_to_s3(file_buffer, filename,metadata)



            return {
                "statusCode": 200,
                "body": json.dumps({
                    "message": "Assignment generated and uploaded successfully.",
                    "fileUrl": file_url,
                    "filename": filename,
                    "questions" : clean_data
                })
            }
            
        
        except Exception as e:
            return {"statusCode": 500, "body": json.dumps({"error": "Unexpected error", "details": str(e)})}
